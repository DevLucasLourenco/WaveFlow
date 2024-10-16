import io
import os
import time
import docx
import zipfile

from copy import deepcopy

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from service.data.archive import Archive



def deltaTime(method):
    def wrapper(self, *args, **kwargs):
        initialTime = time.time()
        mtd = method(self, *args, **kwargs)
        self._timeToGenerate = time.time() - initialTime
        return mtd
    return wrapper



class Builder:
    
    def __init__(self, archive:Archive, baseDocx) -> None:
        self._timeToGenerate:int
        self.__archive = archive
        self.__baseDocx = docx.Document(baseDocx)
        self.__allKeys :str = list(self.__archive.getMetaData()['columns'])
        self.__firstKey:str = self.__allKeys[0]
        self.__listageStringBuilt:list=list()

            
    def __replaceInfosAtDoc(self, records):
        doc_base = deepcopy(self.__baseDocx)
        self.__paragraph(records, doc_base)
        self.__table(records, doc_base)
        return doc_base

                            
    def __paragraph(self, records:dict, doc_base:docx.Document):        
        for para in doc_base.paragraphs:
            for key, value in records.items():
                if key in para.text:
                    for run in para.runs:
                        if key in run.text:
                            
                            primaryKey = self.__getPrimaryKeyFromKeyWDelimiter(key)
                            run.text = run.text.replace(key, str(value))
                            
                            additional_params = self.__archive.getData()[primaryKey]['additional_parameters']
                            
                            if additional_params['font']:
                                run.font.name = additional_params['font']
                                
                            if additional_params['size'] != 0:
                                run.font.size = Pt(additional_params['size'])
                                
                            if additional_params['bold']:
                                run.bold = True
                                
                            if additional_params['italic']:
                                run.italic = True
                
    
    def __table(self, records:dict, doc_base:docx.Document):
        for table in doc_base.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in records.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    primaryKey = self.__getPrimaryKeyFromKeyWDelimiter(key)
                                    additional_params = self.__archive.getData()[primaryKey]['additional_parameters']
                                    
                                    if additional_params['font']:
                                        run.font.name = additional_params['font']
                                        
                                    if additional_params['size'] != 0:
                                        run.font.size = Pt(additional_params['size'])
                                        
                                    if additional_params['bold']:
                                        run.bold = True
                                        
                                    if additional_params['italic']:
                                        run.italic = True

                                                
    def __getPrimaryKeyFromKeyWDelimiter(self, KeyWDelimiter:str):
            return KeyWDelimiter.strip(self.__archive.getDelimiter())
    
                        
    def __getRecordsFromSameIndex(self, index) -> dict:
        d_aux=dict()
        for keyHeader in self.__allKeys:
            information = self.__archive.getData()[keyHeader]['data_column'][index]
            KeyWithDelimiter = self.__archive.getData()[keyHeader]['key_w/Delimiter']
            d_aux.update({KeyWithDelimiter:information})
        return d_aux


    def _zipFiles(self):
        if self.__archive.getFilesGenerated():
            with zipfile.ZipFile(self.__directoryToAlocateFiles + ".zip", 'w') as zipf:
                for i, file in enumerate(self.__archive.getFilesGenerated()):
                    name = self.__listageStringBuilt[i]
                    fileInMemory = self.__addToMemoryBuffer(file)
                    zipf.writestr(name, fileInMemory.getvalue())
            
    
    def __addToMemoryBuffer(self, file):
        fileInMemory = io.BytesIO()
        file.save(fileInMemory)
        fileInMemory.seek(0)
        return fileInMemory
        
    
    def __verifyPossibilityOfDir(self, txt, localSave):
            self.__directoryToAlocateFiles = os.path.dirname(txt)
            if not os.path.exists(self.__directoryToAlocateFiles):
                try:
                    if localSave:
                        os.makedirs(self.__directoryToAlocateFiles)
                except FileNotFoundError as e:
                    pass


    def __getValuesToStringBuilder(self, keys, index):
        listage = list()
        for k in keys:
            listage.append(self.__archive.getData()[k]['data_column'][index])
        return listage
    
    
    def __executeStringBuilder(self, KEYCOLUMN, INDEX, TEXTATFILE):
        valuesToIncrease = self. __getValuesToStringBuilder(KEYCOLUMN, INDEX)
        stringBuild = TEXTATFILE.format(*valuesToIncrease)+'.docx' if not ".docx" in TEXTATFILE else TEXTATFILE.format(*valuesToIncrease)
        self.__listageStringBuilt.append(stringBuild)
        return stringBuild
    
      
    @deltaTime
    def generate(self):
        self.__archive.getFilesGenerated().clear()
        for i, _ in enumerate(self.__archive.getData()[self.__firstKey]['data_column']):
            allRecordsFromIndex = self.__getRecordsFromSameIndex(i)
            doc = self.__replaceInfosAtDoc(allRecordsFromIndex)
            self.__archive.getFilesGenerated().append(doc)
        
        
    def saveAs(self, textAtFile:str, keyColumn:list[str]=[], ZipFile=False, saveLocally=True):
        """You can easly instruct a format string.\n\ne.g.:
        build = Builder(handler.getArchive(), r'example.docx')\n\n
        build.generate()\n\n
        build.saveAs(textAtFile='DOCS/{} - Example How-To - {}', keyColumn=['DATA', 'NOME'], ZipFile=True, saveLocally=True)"""
        
        if self.__archive.getFilesGenerated():
            self.__listageStringBuilt.clear()
            for i, doc in enumerate(self.__archive.getFilesGenerated()):
                stringBuilder = self.__executeStringBuilder(keyColumn, i, textAtFile)
                self.__verifyPossibilityOfDir(stringBuilder, saveLocally)
                
                if saveLocally:
                    doc.save(stringBuilder)
            
            if ZipFile:
                self._zipFiles()
            
                
        else:
            raise RuntimeError('There is no document generated. Should not you generate first?')
    
    
    def getTimeToGenerate(self):
        if self._timeToGenerate:
            return self._timeToGenerate
        
        